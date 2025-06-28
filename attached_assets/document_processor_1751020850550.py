import os
import logging
import PyPDF2
import docx
from typing import Optional

class DocumentProcessor:
    """Handles processing of different document types (TXT, PDF, DOCX)"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def process_document(self, file_path: str) -> Optional[str]:
        """
        Process a document and extract its text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content or None if processing fails
        """
        try:
            if not os.path.exists(file_path):
                self.logger.error(f"File not found: {file_path}")
                return None
            
            file_extension = file_path.lower().split('.')[-1]
            
            if file_extension == 'txt':
                return self._process_txt(file_path)
            elif file_extension == 'pdf':
                return self._process_pdf(file_path)
            elif file_extension == 'docx':
                return self._process_docx(file_path)
            else:
                self.logger.error(f"Unsupported file type: {file_extension}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {str(e)}")
            return None
    
    def _process_txt(self, file_path: str) -> Optional[str]:
        """Process TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content.strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                return content.strip()
            except Exception as e:
                self.logger.error(f"Error reading TXT file {file_path}: {str(e)}")
                return None
        except Exception as e:
            self.logger.error(f"Error processing TXT file {file_path}: {str(e)}")
            return None
    
    def _process_pdf(self, file_path: str) -> Optional[str]:
        """Process PDF files"""
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        content.append(text)
            
            if content:
                return '\n\n'.join(content).strip()
            else:
                self.logger.warning(f"No text content extracted from PDF: {file_path}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            return None
    
    def _process_docx(self, file_path: str) -> Optional[str]:
        """Process DOCX files"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content.append(' | '.join(row_text))
            
            if content:
                return '\n\n'.join(content).strip()
            else:
                self.logger.warning(f"No text content extracted from DOCX: {file_path}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            return None
