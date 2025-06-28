import os
import logging
import PyPDF2
import docx
from typing import Optional

class DocumentProcessor:
    """Handles processing of different document types (TXT, PDF, DOCX)"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def process_document(self, file_path: str) -> Optional[str]:
        """
        Process a document and extract its text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content or None if processing fails
        """
        try:
            if not os.path.exists(file_path):
                self.logger.error(f"File not found: {file_path}")
                return None
            
            # Check if file is empty
            if os.path.getsize(file_path) == 0:
                self.logger.error(f"File is empty: {file_path}")
                return None
            
            file_extension = file_path.lower().split('.')[-1]
            
            if file_extension == 'txt':
                return self._process_txt(file_path)
            elif file_extension == 'pdf':
                return self._process_pdf(file_path)
            elif file_extension == 'docx':
                return self._process_docx(file_path)
            else:
                self.logger.error(f"Unsupported file type: {file_extension}")
                return None
                
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {str(e)}")
            return None
    
    def _process_txt(self, file_path: str) -> Optional[str]:
        """Process TXT files with better encoding handling"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                if content.strip():
                    self.logger.info(f"Successfully read TXT file with {encoding} encoding")
                    return content.strip()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                self.logger.error(f"Error reading TXT file with {encoding}: {str(e)}")
                continue
        
        self.logger.error(f"Could not read TXT file {file_path} with any encoding")
        return None
    
    def _process_pdf(self, file_path: str) -> Optional[str]:
        """Process PDF files with improved error handling"""
        try:
            content = []
            
            with open(file_path, 'rb') as file:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    if len(pdf_reader.pages) == 0:
                        self.logger.warning(f"PDF has no pages: {file_path}")
                        return None
                    
                    for page_num in range(len(pdf_reader.pages)):
                        try:
                            page = pdf_reader.pages[page_num]
                            text = page.extract_text()
                            if text and text.strip():
                                content.append(text.strip())
                        except Exception as e:
                            self.logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                            continue
                    
                except PyPDF2.PdfReadError as e:
                    self.logger.error(f"PDF read error: {str(e)}")
                    return None
                except Exception as e:
                    self.logger.error(f"Unexpected PDF error: {str(e)}")
                    return None
            
            if content:
                combined_content = '\n\n'.join(content).strip()
                if combined_content:
                    self.logger.info(f"Successfully extracted text from PDF with {len(content)} pages")
                    return combined_content
            
            self.logger.warning(f"No text content extracted from PDF: {file_path}")
            return None
                
        except Exception as e:
            self.logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            return None
    
    def _process_docx(self, file_path: str) -> Optional[str]:
        """Process DOCX files with comprehensive content extraction"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_content.append(' | '.join(row_text))
                
                if table_content:
                    content.extend(table_content)
            
            # Extract text from headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            content.append(f"[Header: {text}]")
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            content.append(f"[Footer: {text}]")
            
            if content:
                combined_content = '\n\n'.join(content).strip()
                if combined_content:
                    self.logger.info(f"Successfully extracted text from DOCX with {len(content)} elements")
                    return combined_content
            
            self.logger.warning(f"No text content extracted from DOCX: {file_path}")
            return None
                
        except Exception as e:
            self.logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            return None
